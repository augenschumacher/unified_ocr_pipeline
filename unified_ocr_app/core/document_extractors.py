"""Text extraction helpers for office document inputs."""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path


logger = logging.getLogger("UnifiedOCR")


def _log(log_callback, message: str) -> None:
    if log_callback:
        log_callback(message)


def extract_text_from_docx(docx_path: Path, log_callback=None) -> str:
    """Extract paragraphs and tables from a DOCX file."""
    try:
        import docx

        doc = docx.Document(str(docx_path))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n\n".join(full_text)
    except Exception as e:
        _log(log_callback, f"Fehler beim Extrahieren von Text aus DOCX: {e}")
        logger.exception("Text-Extraktion aus DOCX fehlgeschlagen")
        return ""


def extract_text_from_odt(odt_path: Path, log_callback=None) -> str:
    """Extract paragraphs, headings, and tables from an ODT file."""
    try:
        import zipfile
        import xml.etree.ElementTree as ET

        text_content = []
        with zipfile.ZipFile(odt_path) as z:
            content_xml = z.read("content.xml")
            root = ET.fromstring(content_xml)

            def walk(element):
                local_name = element.tag.split("}")[-1]
                if local_name in ("p", "h"):
                    p_text = "".join(element.itertext()).strip()
                    if p_text:
                        text_content.append(p_text)
                elif local_name == "table":
                    rows = []
                    for child in element:
                        child_local = child.tag.split("}")[-1]
                        if child_local == "table-row":
                            row_cells = []
                            for cell in child:
                                cell_local = cell.tag.split("}")[-1]
                                if cell_local == "table-cell":
                                    row_cells.append("".join(cell.itertext()).strip())
                            if any(row_cells):
                                rows.append(" | ".join(row_cells))
                    if rows:
                        text_content.append("\n".join(rows))
                else:
                    for child in element:
                        if local_name != "table":
                            walk(child)

            walk(root)
        return "\n\n".join(text_content)
    except Exception as e:
        _log(log_callback, f"Fehler beim Extrahieren von Text aus ODT: {e}")
        logger.exception("Text-Extraktion aus ODT fehlgeschlagen")
        return ""


def extract_text_from_doc(doc_path: Path, log_callback=None) -> str:
    """Extract text from legacy Word DOC files, using COM when available."""
    try:
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(doc_path.resolve()))
            text = doc.Content.Text
            doc.Close()
            return text
        finally:
            word.Quit()
    except Exception as e:
        _log(log_callback, f"win32com Extraktion für DOC fehlgeschlagen ({e}). Nutze Fallback...")
        try:
            content = doc_path.read_bytes()
            text_parts = []

            pattern = re.compile(b"(?:[\x20-\x7E\xA0-\xFF\x09\x0A\x0D]\x00){4,}")
            for match in pattern.finditer(content):
                try:
                    decoded = match.group(0).decode("utf-16le").strip()
                    if len(decoded) > 3 and not any(k in decoded for k in ["Normal.dotm", "Microsoft Word", "Title"]):
                        text_parts.append(decoded)
                except Exception:
                    pass
            if text_parts:
                return "\n\n".join(text_parts)

            pattern_ascii = re.compile(b"[\x20-\x7E\x09\x0A\x0D]{4,}")
            for match in pattern_ascii.finditer(content):
                try:
                    decoded = match.group(0).decode("ascii").strip()
                    if len(decoded) > 5 and not any(k in decoded for k in ["Microsoft", "WordDocument", "SummaryInformation"]):
                        text_parts.append(decoded)
                except Exception:
                    pass
            return "\n\n".join(text_parts) if text_parts else "[Kein Text in DOC-Datei gefunden]"
        except Exception as ex:
            _log(log_callback, f"Fallback-Extraktion für DOC fehlgeschlagen: {ex}")
            return ""


def extract_text_from_odoc(odoc_path: Path, log_callback=None) -> str:
    """Extract title and link data from a Synology Office .odoc pointer file."""
    try:
        content = odoc_path.read_text(encoding="utf-8").strip()
        data = json.loads(content)
        title = data.get("title", odoc_path.stem)
        url = data.get("url", "")
        doc_id = data.get("doc_id", "")

        text_lines = [
            "Synology Office Dokument-Link",
            f"Titel: {title}",
        ]
        if url:
            text_lines.append(f"URL: {url}")
        if doc_id:
            text_lines.append(f"Dokument-ID: {doc_id}")
        return "\n".join(text_lines)
    except Exception as e:
        _log(log_callback, f"Konnte .odoc nicht als JSON lesen ({e}), versuche Plaintext...")
        try:
            text = odoc_path.read_text(encoding="utf-8", errors="ignore")
            return f"Synology .odoc Inhalt:\n{text}"
        except Exception as ex:
            _log(log_callback, f"Fehler beim Extrahieren von Text aus .odoc: {ex}")
            return ""
