import re
import logging
from pathlib import Path

logger = logging.getLogger("UnifiedOCR")

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    docx = None

def set_cell_background(cell, fill_color):
    """Setzt die Hintergrundfarbe einer Zelle. fill_color ist ein Hex-String, z.B. '1B365D'."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Setzt das Innenabstand (Padding) einer Zelle in dxa (1/20 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_formatted_text(paragraph, text, is_header=False):
    """Hilfsfunktion für Inline-Formatierungen (Fett und Kursiv)"""
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            if is_header:
                run.font.color.rgb = RGBColor(255, 255, 255)
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            run = paragraph.add_run(token)
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

def render_markdown_to_doc(doc, md_text: str):
    """Rendert Markdown-Text in ein bestehendes docx.Document-Objekt."""
    # Bereinigung: <table_block> und </table_block> XML-Tags entfernen, falls vorhanden
    md_text = md_text.replace("<table_block>", "").replace("</table_block>", "")
    
    # Design-Farben
    COLOR_PRIMARY = RGBColor(27, 54, 93)      # Edles Dunkelblau
    HEX_HEADER_BG = "1B365D"
    HEX_ROW_ALT_BG = "F4F6F9"
    
    lines = md_text.split('\n')
    i = 0
    n = len(lines)
    
    while i < n:
        line = lines[i].strip()
        
        # 1. Tabellen
        if line.startswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            parsed_rows = []
            for tl in table_lines:
                if re.match(r'^\|\s*[-:]+\s*\|', tl) or (tl.count('|') > 1 and all(c in ' |-:' for c in tl)):
                    continue
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                parsed_rows.append(cells)
                
            if parsed_rows:
                num_cols = max(len(row) for row in parsed_rows)
                num_rows = len(parsed_rows)
                
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                
                col_widths = [1.0] * num_cols
                
                for r_idx, row_data in enumerate(parsed_rows):
                    row = table.rows[r_idx]
                    is_header = (r_idx == 0)
                    bg_color = HEX_HEADER_BG if is_header else (HEX_ROW_ALT_BG if r_idx % 2 == 1 else "FFFFFF")
                    
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            set_cell_background(cell, bg_color)
                            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                            p = cell.paragraphs[0]
                            add_formatted_text(p, cell_value, is_header=is_header)
                            
                            val_len = len(cell_value)
                            if val_len > col_widths[c_idx]:
                                col_widths[c_idx] = val_len
                                
                total_len = sum(col_widths)
                if total_len > 0:
                    for c_idx in range(num_cols):
                        width_in_inches = max(0.5, (col_widths[c_idx] / total_len) * 6.0)
                        for row in table.rows:
                            if c_idx < len(row.cells):
                                row.cells[c_idx].width = docx.shared.Inches(width_in_inches)
            continue
            
        # 2. Leerzeilen
        if not line:
            doc.add_paragraph()
            i += 1
            continue
            
        # 3. Überschriften (H1 - H6)
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            add_formatted_text(p, text)
            
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_PRIMARY
                r.font.size = Pt(22 - level * 2)
            i += 1
            continue
            
        # 4. Ungeordnete Listen
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, line[2:])
            i += 1
            continue
            
        # 5. Geordnete Listen
        ordered_list_match = re.match(r'^(\d+)\.\s+(.*)', line)
        if ordered_list_match:
            rest = ordered_list_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, rest)
            i += 1
            continue
            
        # 6. Normaler Fließtext
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_formatted_text(p, line)
        i += 1

def save_markdown_as_docx(text_input, output_path: Path, mode: str = "Lesbare DOCX", image_paths: list = None, quality_report: dict = None) -> Path:
    if not docx:
        logger.error("python-docx fehlt. DOCX Export abgebrochen.")
        raise ImportError("Das Python-Paket 'python-docx' ist nicht installiert. Bitte führen Sie 'pip install python-docx' aus.")

    doc = docx.Document()
    
    # Standard-Schriftart einstellen
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Dictionary zu String konvertieren, falls nötig
    fused_pages = {}
    if isinstance(text_input, dict):
        fused_pages = text_input
        fused_text = "\n\n".join(fused_pages.values())
    else:
        fused_text = str(text_input)
        # Sehr einfache Heuristik zum Splitten nach Seiten (falls String übergeben wurde)
        fused_pages = {1: fused_text}

    if mode == "Prüf-DOCX":
        # Title of Proof Document
        p_title = doc.add_paragraph()
        p_title_run = p_title.add_run("PRÜF-DOKUMENT (Qualitätskontrolle)")
        p_title_run.bold = True
        p_title_run.font.size = Pt(18)
        p_title_run.font.color.rgb = RGBColor(197, 90, 17) # Orange
        
        # Qualitätsbericht am Anfang ausgeben
        if quality_report and quality_report.get("warnings"):
            p_warn_title = doc.add_paragraph()
            p_warn_title.add_run("WARNUNGEN ZUR QUALITÄTSKONTROLLE:").bold = True
            p_warn_title.runs[0].font.color.rgb = RGBColor(197, 0, 0) # Rot
            
            for warning in quality_report["warnings"]:
                p_warn = doc.add_paragraph(style='List Bullet')
                r_warn = p_warn.add_run(warning)
                r_warn.font.color.rgb = RGBColor(197, 0, 0)
            
            # Trennlinie hinzufügen
            doc.add_paragraph().add_run("―" * 40).font.color.rgb = RGBColor(180, 180, 180)

        # Seiten weise rendern
        total_pages = max(fused_pages.keys()) if fused_pages else 0
        for page_num in sorted(fused_pages.keys()):
            # Überschrift für Seite
            p_head = doc.add_paragraph()
            p_head_run = p_head.add_run(f"SEITE {page_num}")
            p_head_run.bold = True
            p_head_run.font.size = Pt(14)
            p_head_run.font.color.rgb = RGBColor(27, 54, 93)
            
            # Bild einbetten
            if image_paths and (page_num - 1) < len(image_paths):
                img_path = image_paths[page_num - 1]
                if Path(img_path).exists():
                    try:
                        p_img = doc.add_paragraph()
                        p_img.add_run().add_picture(str(img_path), width=docx.shared.Inches(5))
                        logger.info(f"Bild {img_path} in Prüf-DOCX eingebettet.")
                    except Exception as ie:
                        logger.error(f"Fehler beim Einbetten des Bildes {img_path}: {ie}")
                        p_err = doc.add_paragraph()
                        p_err.add_run(f"[Bildfehler: {ie}]").italic = True
                else:
                    logger.warning(f"Bildpfad {img_path} existiert nicht.")
            
            # Text darunter rendern
            page_text = fused_pages.get(page_num, "")
            doc.add_paragraph().add_run("Erkannter & korrigierter Text:").bold = True
            render_markdown_to_doc(doc, page_text)
            
            # Seitenumbruch (außer bei der letzten Seite)
            if page_num < total_pages:
                doc.add_page_break()
                
    elif mode == "Originalgetreue DOCX":
        # In diesem Modus liegt der Fokus auf präzisem Tabellenlayout und exakter Rekonstruktion
        p_title = doc.add_paragraph()
        p_title_run = p_title.add_run("Originalgetreues Transkript (Tabellen-Layout)")
        p_title_run.bold = True
        p_title_run.font.size = Pt(14)
        p_title_run.font.color.rgb = RGBColor(27, 54, 93)
        
        render_markdown_to_doc(doc, fused_text)
        
    else: # "Lesbare DOCX" (Standard)
        render_markdown_to_doc(doc, fused_text)

    # Wenn Warnungen vorliegen, hängen wir sie in jedem Modus (außer Prüf-DOCX, dort stehen sie ja schon ganz oben) am Ende des Dokuments an
    if mode != "Prüf-DOCX" and quality_report and quality_report.get("warnings"):
        doc.add_page_break()
        p_warn_title = doc.add_paragraph()
        p_warn_title.add_run("QUALITÄTSHINWEIS (Automatische Prüfung):").bold = True
        p_warn_title.runs[0].font.color.rgb = RGBColor(197, 90, 17) # Orange/Braun
        
        p_desc = doc.add_paragraph()
        p_desc.add_run(
            "Die automatisierte Qualitätsprüfung hat Abweichungen zwischen dem erkannten Text und dem Original festgestellt. "
            "Bitte prüfen Sie folgende Werte manuell im Dokument:"
        ).italic = True
        
        for warning in quality_report["warnings"]:
            p_warn = doc.add_paragraph(style='List Bullet')
            r_warn = p_warn.add_run(warning)
            r_warn.font.color.rgb = RGBColor(197, 90, 17)

    doc.save(str(output_path))
    return output_path
